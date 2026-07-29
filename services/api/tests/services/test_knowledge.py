import io
from collections.abc import Sequence
from datetime import UTC, date, datetime

import pytest
from docx import Document as DocxDocument
from sqlalchemy import event

from app.db.base import Base
from app.db.session import create_database_engine, create_session_factory
from app.models.entities import User
from app.schemas.knowledge import (
    DocumentChunk,
    DocumentFileType,
    DocumentMetadata,
    DocumentRecord,
    KnowledgeCitation,
)
from app.services.knowledge import (
    DuplicateDocumentError,
    GroundedAnswer,
    InMemoryKnowledgeRepository,
    KnowledgePersistenceError,
    KnowledgeService,
    SqlAlchemyKnowledgeRepository,
)
from app.services.knowledge import parser as document_parser


def test_document_title_is_normalized_before_length_validation() -> None:
    normalized_title = "x" * 500

    metadata = DocumentMetadata(
        title=f"  {normalized_title}  ",
        file_type=DocumentFileType.TXT,
    )

    assert metadata.title == normalized_title
    with pytest.raises(ValueError):
        DocumentMetadata(
            title=f"  {normalized_title}x  ",
            file_type=DocumentFileType.TXT,
        )


@pytest.mark.asyncio
async def test_txt_ingest_search_and_evidence_only_answer() -> None:
    service = KnowledgeService(InMemoryKnowledgeRepository())
    record = await service.ingest(
        user_id="user_demo",
        metadata=DocumentMetadata(
            title="人工智能学院考试通知",
            department="人工智能学院",
            publish_date=date(2026, 7, 1),
            applicable_group="2023级",
            version="v1",
            file_type=DocumentFileType.TXT,
        ),
        content="机器学习课程考试时间为2026年7月18日上午九点，地点为教学楼A302。".encode(),
    )

    assert record.chunk_count == 1
    search = await service.search("机器学习考试地点", top_k=5, min_similarity=0.05)
    assert search.results
    assert search.results[0].page_number is None
    assert search.results[0].file_title == "人工智能学院考试通知"

    answer = await service.ask("机器学习考试在哪里", top_k=5, min_similarity=0.05)
    assert answer.sufficient_evidence is True
    assert "教学楼A302" in answer.answer
    assert answer.citations[0].original_text in answer.answer


@pytest.mark.asyncio
async def test_answer_refuses_when_evidence_is_insufficient() -> None:
    service = KnowledgeService(InMemoryKnowledgeRepository())

    answer = await service.ask("奖学金什么时候发放", top_k=3, min_similarity=0.1)

    assert answer.sufficient_evidence is False
    assert answer.citations == []
    assert "没有足够证据" in answer.answer


@pytest.mark.asyncio
async def test_configured_llm_answerer_is_used_only_with_numbered_evidence() -> None:
    class Answerer:
        async def generate(
            self,
            question: str,
            citations: Sequence[KnowledgeCitation],
        ) -> GroundedAnswer:
            assert question == "考试在哪里"
            assert citations[0].original_text.endswith("教学楼A302。")
            return GroundedAnswer(
                answer="考试地点为教学楼 A302。[1]",
                sufficient=True,
                citation_indexes=[1],
            )

    service = KnowledgeService(InMemoryKnowledgeRepository(), answerer=Answerer())
    await service.ingest(
        user_id="user_demo",
        metadata=DocumentMetadata(title="考试通知", file_type=DocumentFileType.TXT),
        content="机器学习考试地点为教学楼A302。".encode(),
    )

    answer = await service.ask("考试在哪里", top_k=3, min_similarity=0.05)

    assert answer.sufficient_evidence is True
    assert answer.answer == "考试地点为教学楼 A302。[1]"
    assert len(answer.citations) == 1


@pytest.mark.asyncio
async def test_uncited_llm_answer_fails_closed_to_original_excerpts() -> None:
    class UncitedAnswerer:
        async def generate(
            self,
            question: str,
            citations: Sequence[KnowledgeCitation],
        ) -> GroundedAnswer:
            del question, citations
            return GroundedAnswer(
                answer="考试地点为教学楼 A302。",
                sufficient=True,
                citation_indexes=[],
            )

    service = KnowledgeService(InMemoryKnowledgeRepository(), answerer=UncitedAnswerer())
    await service.ingest(
        user_id="user_demo",
        metadata=DocumentMetadata(title="考试通知", file_type=DocumentFileType.TXT),
        content="机器学习考试地点为教学楼A302。".encode(),
    )

    answer = await service.ask("考试在哪里", top_k=3, min_similarity=0.05)

    assert answer.sufficient_evidence is True
    assert answer.answer.startswith("根据检索到的校园通知原文")
    assert answer.citations[0].original_text in answer.answer


@pytest.mark.asyncio
async def test_multiple_document_versions_are_reported() -> None:
    service = KnowledgeService(InMemoryKnowledgeRepository())
    for version in ("v1", "v2"):
        await service.ingest(
            user_id="user_demo",
            metadata=DocumentMetadata(
                title="报名通知",
                publish_date=date(2026, 7, 1),
                version=version,
                file_type=DocumentFileType.MARKDOWN,
            ),
            content=f"报名截止时间为7月{17 if version == 'v1' else 18}日。".encode(),
        )

    result = await service.search("报名截止时间", top_k=5, min_similarity=0.05)
    ambiguous = await service.ask("报名截止时间", top_k=5, min_similarity=0.05)
    selected = await service.ask(
        "报名截止时间",
        top_k=5,
        min_similarity=0.05,
        version="v2",
    )

    assert len(result.version_conflicts) == 1
    assert result.version_conflicts[0].versions == ["v1", "v2"]
    assert ambiguous.sufficient_evidence is False
    assert selected.sufficient_evidence is True
    assert len(selected.citations) == 1
    assert selected.citations[0].version == "v2"


@pytest.mark.asyncio
async def test_multiple_applicable_groups_require_clarification_and_can_be_filtered() -> None:
    service = KnowledgeService(InMemoryKnowledgeRepository())
    for group, date_text in (("2023级", "7月17日"), ("2024级", "7月18日")):
        await service.ingest(
            user_id="user_demo",
            metadata=DocumentMetadata(
                title="报名通知",
                applicable_group=group,
                file_type=DocumentFileType.TXT,
            ),
            content=f"{group}报名截止时间为{date_text}。".encode(),
        )

    ambiguous = await service.ask("报名截止时间", top_k=5, min_similarity=0.05)
    filtered = await service.ask(
        "报名截止时间",
        top_k=5,
        min_similarity=0.05,
        applicable_group="2023级学生",
    )

    assert ambiguous.sufficient_evidence is False
    assert ambiguous.applicability_conflicts[0].applicable_groups == ["2023级", "2024级"]
    assert filtered.sufficient_evidence is True
    assert len(filtered.citations) == 1
    assert filtered.citations[0].applicable_group == "2023级"


def test_pdf_page_numbers_come_from_parser_pages(monkeypatch: pytest.MonkeyPatch) -> None:
    class Page:
        def __init__(self, text: str) -> None:
            self.text = text

        def extract_text(self) -> str:
            return self.text

    class Reader:
        is_encrypted = False
        pages = [Page("第一页通知"), Page("第二页通知")]

        def __init__(self, stream: object) -> None:
            del stream

    monkeypatch.setattr(document_parser, "PdfReader", Reader)

    sections = document_parser.parse_document(b"real-parser-input", DocumentFileType.PDF)

    assert [(item.text, item.page_number) for item in sections] == [
        ("第一页通知", 1),
        ("第二页通知", 2),
    ]


def test_docx_parses_paragraphs_and_tables_without_inventing_pages() -> None:
    document = DocxDocument()
    document.add_paragraph("考试安排")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "时间"
    table.cell(0, 1).text = "九点"
    stream = io.BytesIO()
    document.save(stream)

    sections = document_parser.parse_document(stream.getvalue(), DocumentFileType.DOCX)

    assert sections[0].page_number is None
    assert "考试安排" in sections[0].text
    assert "时间\t九点" in sections[0].text


@pytest.mark.asyncio
async def test_sqlalchemy_repository_is_user_scoped_and_rejects_duplicate_content() -> None:
    engine = create_database_engine("sqlite+aiosqlite:///:memory:")
    factory = create_session_factory(engine)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    async with factory() as session, session.begin():
        session.add_all(
            [
                User(id="user_demo", display_name="Demo"),
                User(id="user_other", display_name="Other"),
            ]
        )

    service = KnowledgeService(SqlAlchemyKnowledgeRepository(factory, "user_demo"))
    metadata = DocumentMetadata(title="校历", file_type=DocumentFileType.TXT)
    await service.ingest(user_id="user_demo", metadata=metadata, content="暑假7月20日开始".encode())

    documents = await service.list_documents()
    assert len(documents) == 1
    assert documents[0].chunk_count == 1
    other_service = KnowledgeService(SqlAlchemyKnowledgeRepository(factory, "user_other"))
    assert await other_service.list_documents() == []

    with pytest.raises(DuplicateDocumentError):
        await service.ingest(
            user_id="user_demo",
            metadata=metadata,
            content="暑假7月20日开始".encode(),
        )
    await engine.dispose()


@pytest.mark.asyncio
async def test_nonduplicate_chunk_integrity_error_is_not_reported_as_duplicate() -> None:
    engine = create_database_engine("sqlite+aiosqlite:///:memory:")
    factory = create_session_factory(engine)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    async with factory() as session, session.begin():
        session.add(User(id="user_demo", display_name="Demo"))

    repository = SqlAlchemyKnowledgeRepository(factory, "user_demo")
    record = DocumentRecord(
        id="doc_nonduplicate_integrity",
        user_id="user_demo",
        metadata=DocumentMetadata(
            title="Nonduplicate integrity failure",
            file_type=DocumentFileType.TXT,
        ),
        content_sha256="a" * 64,
        status="ready",
        chunk_count=2,
        created_at=datetime.now(UTC),
    )
    chunks = [
        DocumentChunk(
            id="chunk_nonduplicate_integrity_1",
            document_id=record.id,
            ordinal=0,
            content="first",
        ),
        DocumentChunk(
            id="chunk_nonduplicate_integrity_2",
            document_id=record.id,
            ordinal=0,
            content="second",
        ),
    ]

    with pytest.raises(KnowledgePersistenceError, match="document write failed"):
        await repository.save(record, chunks)
    assert await repository.list_documents() == []
    await engine.dispose()


@pytest.mark.asyncio
async def test_document_list_uses_one_query_for_all_chunk_counts() -> None:
    engine = create_database_engine("sqlite+aiosqlite:///:memory:")
    factory = create_session_factory(engine)
    async with engine.begin() as connection:
        await connection.run_sync(Base.metadata.create_all)
    async with factory() as session, session.begin():
        session.add(User(id="user_demo", display_name="Demo"))

    service = KnowledgeService(SqlAlchemyKnowledgeRepository(factory, "user_demo"))
    for index in range(3):
        await service.ingest(
            user_id="user_demo",
            metadata=DocumentMetadata(
                title=f"Document {index}",
                file_type=DocumentFileType.TXT,
            ),
            content=f"document body {index}".encode(),
        )

    select_statements: list[str] = []

    def count_selects(
        _connection: object,
        _cursor: object,
        statement: str,
        _parameters: object,
        _context: object,
        _executemany: object,
    ) -> None:
        if statement.lstrip().upper().startswith("SELECT"):
            select_statements.append(statement)

    event.listen(engine.sync_engine, "before_cursor_execute", count_selects)
    try:
        documents = await service.list_documents()
    finally:
        event.remove(engine.sync_engine, "before_cursor_execute", count_selects)

    assert len(documents) == 3
    assert [document.chunk_count for document in documents] == [1, 1, 1]
    assert len(select_statements) == 1
    await engine.dispose()
