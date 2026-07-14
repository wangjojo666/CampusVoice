import io
import re
import zipfile
from dataclasses import dataclass

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.schemas.knowledge import DocumentFileType


class DocumentParseError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code
        self.message = message


class DocumentLimitError(DocumentParseError):
    """Raised when a document exceeds a bounded ingestion resource limit."""


@dataclass(frozen=True, slots=True)
class DocumentLimits:
    max_pdf_pages: int = 500
    max_docx_entries: int = 10_000
    max_docx_uncompressed_bytes: int = 100 * 1024 * 1024
    max_docx_compression_ratio: float = 100.0
    max_extracted_characters: int = 1_000_000
    max_chunks: int = 2_000


DEFAULT_DOCUMENT_LIMITS = DocumentLimits()


@dataclass(frozen=True, slots=True)
class ParsedSection:
    text: str
    page_number: int | None


def _clean_text(value: str) -> str:
    value = value.replace("\x00", "")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _check_extracted_character_limit(character_count: int, limits: DocumentLimits) -> None:
    if character_count > limits.max_extracted_characters:
        raise DocumentLimitError(
            "document_text_too_large",
            "文档提取文本不能超过 1,000,000 个字符。",
        )


def _validate_docx_archive(content: bytes, limits: DocumentLimits) -> None:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        entries = archive.infolist()
        if len(entries) > limits.max_docx_entries:
            raise DocumentLimitError(
                "document_archive_entry_limit_exceeded",
                "DOCX ZIP 条目不能超过 10,000 个。",
            )
        total_uncompressed = sum(entry.file_size for entry in entries)
        total_compressed = sum(entry.compress_size for entry in entries)
        if total_uncompressed > limits.max_docx_uncompressed_bytes:
            raise DocumentLimitError(
                "document_expanded_size_too_large",
                "DOCX 解压后的总大小不能超过 100 MB。",
            )
        suspicious_entry = any(
            entry.file_size
            and (
                entry.compress_size == 0
                or entry.file_size / entry.compress_size > limits.max_docx_compression_ratio
            )
            for entry in entries
        )
        suspicious_archive = total_uncompressed and (
            total_compressed == 0
            or total_uncompressed / total_compressed > limits.max_docx_compression_ratio
        )
        if suspicious_entry or suspicious_archive:
            raise DocumentLimitError(
                "suspicious_docx_compression",
                "DOCX 压缩比异常，已拒绝解析。",
            )


def parse_document(
    content: bytes,
    file_type: DocumentFileType,
    *,
    limits: DocumentLimits = DEFAULT_DOCUMENT_LIMITS,
) -> list[ParsedSection]:
    if not content:
        raise DocumentParseError("empty_document", "上传的文档为空。")
    try:
        if file_type == DocumentFileType.PDF:
            reader = PdfReader(io.BytesIO(content))
            if reader.is_encrypted:
                try:
                    unlocked = reader.decrypt("")
                except Exception as exc:
                    raise DocumentParseError(
                        "encrypted_document", "无法读取加密 PDF，请先移除密码。"
                    ) from exc
                if not unlocked:
                    raise DocumentParseError(
                        "encrypted_document", "无法读取加密 PDF，请先移除密码。"
                    )
            if len(reader.pages) > limits.max_pdf_pages:
                raise DocumentLimitError(
                    "document_page_limit_exceeded",
                    "PDF 不能超过 500 页。",
                )
            sections = []
            extracted_characters = 0
            for index, page in enumerate(reader.pages):
                text = _clean_text(page.extract_text() or "")
                extracted_characters += len(text)
                _check_extracted_character_limit(extracted_characters, limits)
                sections.append(ParsedSection(text, page_number=index + 1))
        elif file_type == DocumentFileType.DOCX:
            _validate_docx_archive(content, limits)
            document = DocxDocument(io.BytesIO(content))
            paragraphs = [_clean_text(paragraph.text) for paragraph in document.paragraphs]
            table_rows = [
                "\t".join(_clean_text(cell.text) for cell in row.cells)
                for table in document.tables
                for row in table.rows
            ]
            text = "\n".join(filter(None, [*paragraphs, *table_rows]))
            _check_extracted_character_limit(len(text), limits)
            # DOCX has no reliable page mapping without a layout engine, so page is null.
            sections = [ParsedSection(text, page_number=None)]
        elif file_type in {DocumentFileType.TXT, DocumentFileType.MARKDOWN}:
            try:
                decoded = content.decode("utf-8-sig")
            except UnicodeDecodeError as exc:
                raise DocumentParseError(
                    "unsupported_encoding", "TXT/Markdown 文档必须使用 UTF-8 编码。"
                ) from exc
            text = _clean_text(decoded)
            _check_extracted_character_limit(len(text), limits)
            sections = [ParsedSection(text, page_number=None)]
        else:
            raise DocumentParseError("unsupported_file_type", "不支持的文档类型。")
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError("document_parse_failed", "文档解析失败，请检查文件内容。") from exc

    nonempty = [section for section in sections if section.text]
    if not nonempty:
        raise DocumentParseError(
            "no_extractable_text",
            "文档中没有可提取的文本；扫描版 PDF 需要先进行 OCR。",
        )
    return nonempty


def split_sections(
    sections: list[ParsedSection],
    *,
    chunk_size: int = 700,
    overlap: int = 100,
    max_chunks: int = DEFAULT_DOCUMENT_LIMITS.max_chunks,
) -> list[ParsedSection]:
    if chunk_size <= 0 or overlap < 0 or overlap >= chunk_size or max_chunks <= 0:
        raise ValueError(
            "chunk_size and max_chunks must be positive and overlap smaller than chunk_size"
        )
    chunks: list[ParsedSection] = []
    for section in sections:
        text = section.text
        start = 0
        while start < len(text):
            hard_end = min(len(text), start + chunk_size)
            end = hard_end
            if hard_end < len(text):
                boundary = max(
                    text.rfind("\n", start + chunk_size // 2, hard_end),
                    text.rfind("。", start + chunk_size // 2, hard_end),
                    text.rfind("；", start + chunk_size // 2, hard_end),
                )
                if boundary > start:
                    end = boundary + 1
            piece = text[start:end].strip()
            if piece:
                if len(chunks) >= max_chunks:
                    raise DocumentLimitError(
                        "document_chunk_limit_exceeded",
                        "文档切分后的片段不能超过 2,000 个。",
                    )
                chunks.append(ParsedSection(piece, section.page_number))
            if end >= len(text):
                break
            start = max(start + 1, end - overlap)
    return chunks
